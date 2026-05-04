"""INDForge - preclinical IND-enabling biomarker package app.

Serves static assets from indforge_static/ and a small set of mock JSON
endpoints. Live endpoints are stubs — when wired to Domino, they'd read
from the biomarker-forge-outputs NetApp volume plus tox/PK stores.
"""

from __future__ import annotations

import io
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import FastAPI
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# python-docx is a hard dependency for FIH dose justification export.
from docx import Document
from docx.shared import Pt, RGBColor

# Optional: Domino SDK @add_tracing decorator for the LLM-drafted briefing
# paragraph. Falls back to a no-op when the SDK isn't installed locally.
try:
    from domino.sdk import add_tracing  # type: ignore
except Exception:  # noqa: BLE001
    def add_tracing(**_kwargs: Any):  # type: ignore[no-redef]
        def deco(fn):
            return fn
        return deco

BASE_DIR = Path(__file__).resolve().parent
STATIC_DIR = BASE_DIR / "indforge_static"

app = FastAPI(title="INDForge")
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


@app.get("/")
def root():
    return FileResponse(str(STATIC_DIR / "index.html"))


@app.get("/api/health")
def health():
    return {"status": "ok", "app": "INDForge"}


@app.get("/api/context")
def context():
    return {
        "app": "INDForge",
        "project_id": os.environ.get("DOMINO_PROJECT_ID"),
        "project_name": os.environ.get("DOMINO_PROJECT_NAME"),
        "project_owner": os.environ.get("DOMINO_PROJECT_OWNER"),
    }


@app.get("/api/candidates")
def candidates():
    # Live endpoint stub. In production this reads the scored shortlist
    # produced by BiomarkerForge and re-ranks on preclinical evidence.
    return JSONResponse(
        {"error": "live_data_unavailable",
         "detail": "Connect to biomarker-forge-outputs volume and GLP tox store."},
        status_code=503,
    )


@app.get("/api/safety")
def safety():
    return JSONResponse({"error": "live_data_unavailable"}, status_code=503)


@app.get("/api/pkpd")
def pkpd():
    return JSONResponse({"error": "live_data_unavailable"}, status_code=503)


@app.get("/api/ind-package")
def ind_package():
    return JSONResponse({"error": "live_data_unavailable"}, status_code=503)


# ----- Translatability Sweep (multi-omics integration on SLURM) -----
# Demo endpoints. In production these are gated behind API_GAPS.liveSweep
# and read from the on-prem multi-omics integration outputs (NetApp).

API_GAPS = {
    "liveSweep": False,
    "liveDoseCalc": False,
    "liveDoseDraft": False,
    "liveLineage": False,
}

# Cohorts the user can sweep over. Pinned to on-prem because animal omics
# is restricted by IT policy.
_SWEEP_COHORTS = [
    {"id": "NHP-2026-A1", "label": "NHP-2026-A1 · 28-day repeat-dose",
     "n_animals": 40, "glp_study": "GLP-002", "dose_groups": 4, "timepoints": 2, "species": "NHP"},
    {"id": "RAT-2026-B2", "label": "RAT-2026-B2 · 14-day exploratory",
     "n_animals": 24, "glp_study": "GLP-001", "dose_groups": 3, "timepoints": 2, "species": "Rat"},
]


@app.get("/api/sweep/cohorts")
def sweep_cohorts():
    return {"cohorts": _SWEEP_COHORTS}


@app.post("/api/sweep/submit")
def sweep_submit():
    if API_GAPS["liveSweep"]:
        return JSONResponse({"error": "live_data_unavailable"}, status_code=503)
    return {
        "sweep_id": "SWEEP-2026-05-03-A",
        "status": "queued",
        "eta_seconds": 6,
        "partition": "cpu-long",
    }


@app.get("/api/sweep/{sweep_id}/status")
def sweep_status(sweep_id: str):
    if API_GAPS["liveSweep"]:
        return JSONResponse({"error": "live_data_unavailable"}, status_code=503)
    return {"sweep_id": sweep_id, "status": "succeeded", "progress": 1.0,
            "partition": "cpu-long", "started_at": "2026-05-03T14:11:27Z",
            "eta_seconds": 0}


@app.get("/api/sweep/{sweep_id}/results")
def sweep_results(sweep_id: str):
    if API_GAPS["liveSweep"]:
        return JSONResponse({"error": "live_data_unavailable"}, status_code=503)
    # Frontend mock holds the full results payload (factor weights, candidates,
    # audit). Live endpoint would assemble from on-prem outputs + audit store.
    return JSONResponse({"error": "live_data_unavailable", "sweep_id": sweep_id}, status_code=503)


# =====================================================================
# FIH dose calculator + dose justification document export
# =====================================================================
#
# The dose math is reviewable arithmetic, not a model. The same formulas
# run client-side (dose_math.js) so the UI updates live; this server
# implementation is the canonical reference used for the .docx export
# audit page. If the two ever diverge, the server wins — that's the
# trace-of-record.
#
# Citations baked in:
#   - FDA 2005, "Estimating the Maximum Safe Starting Dose in Initial
#     Clinical Trials for Therapeutics in Adult Healthy Volunteers"
#   - EMA 2017, "Guideline on Strategies to Identify and Mitigate Risks
#     for First-in-Human and Early Clinical Trials"
#   - ICH M3(R2), "Nonclinical Safety Studies for the Conduct of Human
#     Clinical Trials and Marketing Authorization for Pharmaceuticals"
#   - ICH S9 (oncology only), "Nonclinical Evaluation for Anticancer
#     Pharmaceuticals"

DOSE_MATH_KERNEL_VERSION = "indforge.dose_math/1.0.0"


class DoseInputs(BaseModel):
    species: str = "NHP"
    animal_bw_kg: float = 5.0
    human_bw_kg: float = 60.0
    noael_mg_kg: float = 6.0
    glp_study_id: Optional[str] = "GLP-002"
    exponent: float = 0.67
    safety_factor: float = 10.0
    modality: str = "biologic"  # "small_molecule" | "biologic"
    # Biologic-only MABEL inputs
    kd_nm: Optional[float] = 12.0
    mw_kda: Optional[float] = 150.0
    vd_l: Optional[float] = 4.0
    ro_threshold: Optional[float] = 0.10
    # Exposure-margin context (optional; sourced from MOCK_PKPD)
    noael_auc: Optional[float] = 3800.0
    auc_per_mg: Optional[float] = 510.0  # projected AUC at 1 mg starting dose


def _round(x: float, n: int = 4) -> float:
    return round(x, n)


def compute_dose(inp: DoseInputs) -> Dict[str, Any]:
    """Reference FIH dose computation. Mirrors dose_math.js exactly."""
    a = inp.animal_bw_kg
    h = inp.human_bw_kg
    noael = inp.noael_mg_kg
    sf = inp.safety_factor

    # HED via allometric scaling. FDA 2005 default exponent is 0.67
    # (surface-area scaling); Boxenbaum BW^0.75 is an alternative.
    scaling = (a / h) ** (1.0 - inp.exponent)
    hed_mg_per_kg = noael * scaling
    mrsd_mg_per_kg = hed_mg_per_kg / sf
    mrsd_mg = mrsd_mg_per_kg * h

    formulas: List[Dict[str, Any]] = [
        {
            "id": "scaling",
            "label": "Allometric scaling factor",
            "formula": f"({a} / {h}) ^ (1 − {inp.exponent})",
            "value": _round(scaling),
            "citation": "FDA 2005 §3 (BW^0.67 surface-area scaling)",
        },
        {
            "id": "hed",
            "label": "HED (Human Equivalent Dose)",
            "formula": f"{noael} mg/kg × {_round(scaling)} = {_round(hed_mg_per_kg)} mg/kg",
            "value": _round(hed_mg_per_kg),
            "unit": "mg/kg",
            "citation": "FDA 2005 §4",
        },
        {
            "id": "mrsd",
            "label": "MRSD (Maximum Recommended Starting Dose)",
            "formula": f"{_round(hed_mg_per_kg)} mg/kg / {sf} = {_round(mrsd_mg_per_kg)} mg/kg",
            "value": _round(mrsd_mg_per_kg),
            "unit": "mg/kg",
            "citation": "FDA 2005 §5 (default safety factor 10)",
        },
        {
            "id": "mrsd_total",
            "label": "MRSD total dose",
            "formula": f"{_round(mrsd_mg_per_kg)} mg/kg × {h} kg = {_round(mrsd_mg)} mg",
            "value": _round(mrsd_mg),
            "unit": "mg",
        },
    ]

    mabel_dose_mg: Optional[float] = None
    if inp.modality == "biologic" and inp.kd_nm and inp.mw_kda and inp.vd_l and inp.ro_threshold:
        # Simplified one-site receptor occupancy model.
        # RO = C / (Kd + C)  →  C = Kd * RO / (1 - RO)
        ro = inp.ro_threshold
        kd = inp.kd_nm
        c_nm = kd * ro / (1.0 - ro)
        # nM → mg/L: nM × MW(g/mol) × 1e-6  (since 1 nM = 1e-9 mol/L)
        mw_g_per_mol = inp.mw_kda * 1000.0
        c_mg_per_l = c_nm * mw_g_per_mol * 1e-6
        mabel_dose_mg = c_mg_per_l * inp.vd_l
        formulas.append({
            "id": "mabel_conc",
            "label": "Target free drug concentration at 10% RO",
            "formula": f"{kd} nM × {ro} / (1 − {ro}) = {_round(c_nm)} nM",
            "value": _round(c_nm),
            "unit": "nM",
            "citation": "EMA 2017 §4.4 (MABEL anchored at 10% RO)",
        })
        formulas.append({
            "id": "mabel_dose",
            "label": "MABEL dose",
            "formula": f"{_round(c_mg_per_l)} mg/L × {inp.vd_l} L = {_round(mabel_dose_mg)} mg",
            "value": _round(mabel_dose_mg),
            "unit": "mg",
            "citation": "EMA 2017 §4.4 (MABEL preferred for biologics)",
        })

    # Controlling method = lower of MRSD-total and MABEL.
    candidates = [("MRSD", mrsd_mg)]
    if mabel_dose_mg is not None:
        candidates.append(("MABEL", mabel_dose_mg))
    controlling_method, fih_dose_mg = min(candidates, key=lambda t: t[1])

    # Exposure margin vs NOAEL: NOAEL_AUC / projected_AUC_at_starting_dose.
    exposure_margin: Optional[float] = None
    if inp.noael_auc and inp.auc_per_mg and fih_dose_mg > 0:
        projected_auc = inp.auc_per_mg * fih_dose_mg
        if projected_auc > 0:
            exposure_margin = inp.noael_auc / projected_auc
            formulas.append({
                "id": "exposure_margin",
                "label": "Exposure margin vs NOAEL",
                "formula": (f"NOAEL AUC {inp.noael_auc} ÷ projected AUC at "
                            f"{_round(fih_dose_mg, 2)} mg ({_round(projected_auc)} ng·h/mL) "
                            f"= {_round(exposure_margin, 1)}×"),
                "value": _round(exposure_margin, 2),
                "unit": "× vs NOAEL",
                "citation": "ICH M3(R2) §3.1",
            })

    rationale_parts = []
    if controlling_method == "MABEL":
        rationale_parts.append(
            "MABEL-controlled: anchored at 10% receptor occupancy per EMA 2017 §4.4 "
            "(MABEL preferred for biologics post-TGN1412)."
        )
    else:
        rationale_parts.append(
            "MRSD-controlled: derived from animal NOAEL via allometric scaling "
            "(FDA 2005) with a 10× safety factor."
        )
    if mabel_dose_mg is not None:
        rationale_parts.append(
            f"Cross-check: MRSD = {_round(mrsd_mg, 2)} mg, "
            f"MABEL = {_round(mabel_dose_mg, 2)} mg; lower value selected."
        )
    rationale = " ".join(rationale_parts)

    return {
        "kernel_version": DOSE_MATH_KERNEL_VERSION,
        "inputs": inp.model_dump(),
        "hed_mg_per_kg": _round(hed_mg_per_kg),
        "mrsd_mg_per_kg": _round(mrsd_mg_per_kg),
        "mrsd_mg": _round(mrsd_mg),
        "mabel_dose_mg": _round(mabel_dose_mg) if mabel_dose_mg is not None else None,
        "fih_dose_mg": _round(fih_dose_mg, 2),
        "controlling_method": controlling_method,
        "exposure_margin": _round(exposure_margin, 2) if exposure_margin else None,
        "formulas": formulas,
        "rationale": rationale,
        "citations": [
            "FDA 2005 — Estimating the Maximum Safe Starting Dose",
            "EMA 2017 — FIH Risk Identification and Mitigation Guideline",
            "ICH M3(R2) — Nonclinical Safety Studies",
        ],
    }


@app.post("/api/dose/calculate")
def dose_calculate(inp: DoseInputs):
    if API_GAPS["liveDoseCalc"]:
        return JSONResponse({"error": "live_data_unavailable"}, status_code=503)
    return compute_dose(inp)


# ----- GLP studies (mock) -----
_GLP_STUDIES = [
    {
        "id": "GLP-001", "species": "Rat", "study_type": "28-day repeat-dose GLP",
        "animal_bw_kg": 0.25, "n_animals": 80, "dose_groups_mg_per_kg": [0, 3, 10, 30],
        "noael_mg_kg": 12.0, "study_director": "K. Hartwell, DVM, DABT",
        "study_date": "2025-09-12", "cro": "Charles River Mattawan",
        "key_findings": [
            {"day": 14, "severity": "Mild", "organ": "Liver",
             "text": "Mild ALT elevation at 30 mg/kg (4 of 10 animals); recovered by Day 28; HE histopathology unremarkable."},
        ],
        "audit_chain_hash": "sha256:e2c4…a91f",
    },
    {
        "id": "GLP-002", "species": "NHP", "study_type": "28-day repeat-dose GLP",
        "animal_bw_kg": 5.0, "n_animals": 40, "dose_groups_mg_per_kg": [0, 2, 6, 20],
        "noael_mg_kg": 6.0, "study_director": "M. Okafor, PhD, DABT",
        "study_date": "2025-11-04", "cro": "Labcorp Madison",
        "key_findings": [
            {"day": 14, "severity": "Moderate", "organ": "CRS / Immune",
             "text": "Transient IL-6 spike at 20 mg/kg (Cmax 142 pg/mL); resolved by Day 21; no clinical correlate."},
            {"day": 28, "severity": "Minimal", "organ": "Liver",
             "text": "Minimal ALT trend at 20 mg/kg (within reference range); not adverse."},
        ],
        "audit_chain_hash": "sha256:9bd8…71c2",
    },
    {
        "id": "GLP-003", "species": "Rat", "study_type": "Genotoxicity panel (Ames + in vivo MN)",
        "animal_bw_kg": 0.25, "n_animals": 30, "dose_groups_mg_per_kg": [0, 50, 200, 500],
        "noael_mg_kg": None, "study_director": "S. Chen, MS, DABT",
        "study_date": "2025-10-20", "cro": "Inotiv",
        "key_findings": [
            {"day": 0, "severity": "None", "organ": "Genotox",
             "text": "Negative Ames (5 strains, ±S9). Negative in vivo bone marrow micronucleus."},
        ],
        "audit_chain_hash": "sha256:4af3…d827",
    },
    {
        "id": "GLP-004", "species": "Rat", "study_type": "Cardiovascular safety pharmacology",
        "animal_bw_kg": 0.25, "n_animals": 16, "dose_groups_mg_per_kg": [0, 5, 15],
        "noael_mg_kg": 15.0, "study_director": "P. Liang, PhD",
        "study_date": "2025-08-30", "cro": "Charles River",
        "key_findings": [
            {"day": 0, "severity": "None", "organ": "Cardiac",
             "text": "No QTc prolongation, no cTnI elevation across all dose groups."},
        ],
        "audit_chain_hash": "sha256:7c10…b4e9",
    },
]


@app.get("/api/glp-studies")
def glp_studies():
    return {"studies": _GLP_STUDIES}


@app.get("/api/glp-studies/{study_id}")
def glp_study(study_id: str):
    for s in _GLP_STUDIES:
        if s["id"] == study_id:
            return s
    return JSONResponse({"error": "not_found", "study_id": study_id}, status_code=404)


# ----- Audit lineage -----
@app.get("/api/lineage/{value_id}")
def lineage(value_id: str):
    if API_GAPS["liveLineage"]:
        return JSONResponse({"error": "live_data_unavailable"}, status_code=503)
    # Placeholder: the frontend assembles lineage from the live dose-calc
    # response. This endpoint exists as the canonical hand-off shape for
    # downstream Vault RIM lift.
    now = datetime.now(timezone.utc).isoformat()
    return {
        "value_id": value_id,
        "snapshot_id": f"snap_{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}_indforge",
        "kernel_version": DOSE_MATH_KERNEL_VERSION,
        "user": {"id": os.environ.get("DOMINO_PROJECT_OWNER", "demo.user@sponsor.com"), "role": "Pharmacometrics Lead"},
        "timestamp": now,
        "vault_lift_target": "Veeva Vault RIM · Submissions · Pre-IND",
    }


# ----- Dose justification draft (LLM-drafted briefing paragraph) -----

class DraftRequest(BaseModel):
    compound_id: str = "INDF-127"
    sweep_id: Optional[str] = None
    dose_inputs: Optional[Dict[str, Any]] = None


_DUMMY_PARAGRAPH_INDF127 = (
    "INDF-127 is a humanized monoclonal antibody targeting PD-1 with a Kd of 1.2 nM. "
    "Based on the 28-day GLP repeat-dose study in cynomolgus macaques (GLP-002), the "
    "NOAEL of 6 mg/kg supports a HED of 2.6 mg/kg via FDA 2005 surface-area scaling and "
    "an MRSD of 0.26 mg/kg (16 mg total for a 60 kg subject) after a 10× safety factor. "
    "Per EMA 2017 §4.4, MABEL is the preferred starting-dose method for receptor-binding "
    "biologics: anchored at 10% receptor occupancy with the program's affinity, distribution, "
    "and molecular weight, MABEL yields 0.8 mg total — the controlling value. The proposed "
    "FIH starting dose is 0.8 mg, providing a 9.3× exposure margin vs the NHP NOAEL AUC "
    "(3800 ng·h/mL). The biomarker plan is anchored on CXCL9 (efficacy/PD) with safety "
    "monitoring of ALT, cTnI, CREA, and IL-6, all qualified through the GLP-002 panel."
)


@add_tracing(name="indforge.dose_justification.draft", trace_io=True)
def _draft_paragraph_via_llm(req: DraftRequest) -> Dict[str, Any]:
    """Draft the executive-summary paragraph using Claude with prompt caching.

    The static prefix (FDA/EMA/ICH outline + briefing-book template, ~3000
    tokens) is cached so each draft hits the cache. Only enabled when
    API_GAPS.liveDoseDraft is True; otherwise the dummy paragraph is used.
    """
    try:
        import anthropic  # type: ignore
    except Exception:
        return {"text": _DUMMY_PARAGRAPH_INDF127, "model_id": "dummy", "live": False}

    client = anthropic.Anthropic()
    static_prefix = (
        "You are drafting the executive-summary paragraph of a FIH dose "
        "justification briefing book. Cite FDA 2005 (MRSD), EMA 2017 (FIH/MABEL), "
        "and ICH M3(R2). Keep to ~150 words, professional regulatory tone, no "
        "exclamation points. Show the controlling method (MRSD or MABEL) and the "
        "exposure margin vs NOAEL. Reference any GLP studies by ID."
    )
    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=400,
        system=[
            {"type": "text", "text": static_prefix,
             "cache_control": {"type": "ephemeral"}},
        ],
        messages=[
            {"role": "user", "content": [
                {"type": "text", "text": f"Compound: {req.compound_id}. Inputs: {req.dose_inputs}."}
            ]},
        ],
    )
    text = "".join(b.text for b in msg.content if getattr(b, "type", None) == "text")
    return {"text": text, "model_id": "claude-sonnet-4-6", "live": True,
            "usage": getattr(msg, "usage", None) and msg.usage.model_dump()}


# In-memory store for drafted documents.
_DRAFTS: Dict[str, Dict[str, Any]] = {}


@app.post("/api/dose/justification/draft")
def dose_justification_draft(req: DraftRequest):
    doc_id = "DOSE-JUST-" + datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ") + "-" + uuid.uuid4().hex[:6]
    if API_GAPS["liveDoseDraft"]:
        result = _draft_paragraph_via_llm(req)
    else:
        # Demo path: pre-authored paragraph; LLM call shape is preserved
        # in the trace metadata so the audit page reads identically.
        result = {"text": _DUMMY_PARAGRAPH_INDF127, "model_id": "demo-static",
                  "live": False, "usage": None}

    inputs = req.dose_inputs or {}
    try:
        dose_result = compute_dose(DoseInputs(**inputs)) if inputs else compute_dose(DoseInputs())
    except Exception:
        dose_result = compute_dose(DoseInputs())

    snapshot_id = f"snap_{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}_indforge"
    record = {
        "document_id": doc_id,
        "compound_id": req.compound_id,
        "sweep_id": req.sweep_id,
        "snapshot_id": snapshot_id,
        "paragraph_text": result["text"],
        "model_id": result["model_id"],
        "prompt_version": "fih-justification-v1.0",
        "live": result["live"],
        "drafted_at": datetime.now(timezone.utc).isoformat(),
        "user_id": os.environ.get("DOMINO_PROJECT_OWNER", "demo.user@sponsor.com"),
        "dose_result": dose_result,
    }
    _DRAFTS[doc_id] = record
    return {"document_id": doc_id, "paragraph_text": record["paragraph_text"],
            "model_id": record["model_id"], "snapshot_id": snapshot_id,
            "prompt_version": record["prompt_version"], "live": record["live"]}


@app.get("/api/dose/justification/{document_id}/export.docx")
def dose_justification_export(document_id: str):
    record = _DRAFTS.get(document_id)
    if record is None:
        # Synthesize on-demand for direct export (demo mode).
        record = {
            "document_id": document_id,
            "compound_id": "INDF-127",
            "sweep_id": "SWEEP-2026-05-03-A",
            "snapshot_id": f"snap_{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}_indforge",
            "paragraph_text": _DUMMY_PARAGRAPH_INDF127,
            "model_id": "demo-static",
            "prompt_version": "fih-justification-v1.0",
            "live": False,
            "drafted_at": datetime.now(timezone.utc).isoformat(),
            "user_id": os.environ.get("DOMINO_PROJECT_OWNER", "demo.user@sponsor.com"),
            "dose_result": compute_dose(DoseInputs()),
        }

    buf = io.BytesIO()
    doc = Document()
    styles = doc.styles
    if "Normal" in [s.name for s in styles]:
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(11)

    def h1(text: str):
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x54, 0x3F, 0xDE)

    def h2(text: str):
        doc.add_heading(text, level=2)

    def p(text: str):
        doc.add_paragraph(text)

    def kv(k: str, v: str):
        para = doc.add_paragraph()
        run = para.add_run(f"{k}: ")
        run.bold = True
        para.add_run(v)

    # 1. Cover
    h1(f"FIH Dose Justification — {record['compound_id']}")
    p("Pre-IND briefing-book extract · Pharmacology and Toxicology section")
    kv("Document ID", record["document_id"])
    kv("Compound", record["compound_id"])
    kv("Snapshot ID", record["snapshot_id"])
    kv("Drafted at", record["drafted_at"])
    kv("Drafted by", record["user_id"])
    kv("Prompt version", record["prompt_version"])
    kv("LLM model", record["model_id"])
    kv("LLM live call", "yes" if record["live"] else "no (demo mode — pre-authored copy)")
    kv("Translatability sweep", record.get("sweep_id") or "—")

    # 2. Executive summary
    h2("Executive summary")
    p(record["paragraph_text"])

    dr = record["dose_result"]
    inp = dr["inputs"]

    # 3. Study basis
    h2("Study basis")
    if inp.get("glp_study_id"):
        p(f"Pivotal NOAEL determination: {inp['glp_study_id']} "
          f"({inp['species']} {inp['noael_mg_kg']} mg/kg).")
    p("Supporting studies: GLP-001 (rat 28-day), GLP-003 (genotoxicity panel), "
      "GLP-004 (cardiovascular safety pharmacology).")

    # 4. Allometric scaling math walkthrough
    h2("Allometric scaling and MRSD derivation (FDA 2005)")
    for f in dr["formulas"]:
        if f["id"] in ("scaling", "hed", "mrsd", "mrsd_total"):
            doc.add_paragraph(f"{f['label']}: {f['formula']}", style="Intense Quote")
            if f.get("citation"):
                doc.add_paragraph(f"   Citation: {f['citation']}")

    # 5. MABEL cross-check
    if any(f["id"] == "mabel_dose" for f in dr["formulas"]):
        h2("MABEL cross-check (EMA 2017)")
        for f in dr["formulas"]:
            if f["id"] in ("mabel_conc", "mabel_dose"):
                doc.add_paragraph(f"{f['label']}: {f['formula']}", style="Intense Quote")
                if f.get("citation"):
                    doc.add_paragraph(f"   Citation: {f['citation']}")

    # 6. Recommendation + exposure margin
    h2("FIH dose recommendation")
    kv("Recommended starting dose", f"{dr['fih_dose_mg']} mg")
    kv("Controlling method", dr["controlling_method"])
    if dr.get("exposure_margin"):
        kv("Exposure margin vs NOAEL", f"{dr['exposure_margin']}× (ICH M3(R2) §3.1)")
    p(dr["rationale"])

    # 7. Sensitivity envelope (translatability sweep)
    h2("Sensitivity envelope — bridging evidence")
    if record.get("sweep_id"):
        p(f"Translatability sweep run {record['sweep_id']} provides per-candidate "
          "translatability scores with 95% bootstrap confidence intervals across "
          "RNA-seq, proteomics, and metabolomics. CXCL9 (efficacy/PD): 0.87 [0.81, 0.92]. "
          "Compute: on-prem SLURM (Cambridge HPC), no data egress, GLP audit chain preserved.")
    else:
        p("No sensitivity sweep attached. Pre-IND meeting will reference the static "
          "single-species NOAEL only.")

    # 8. Safety biomarker panel
    h2("Safety biomarker panel")
    p("ALT, cTnI, CREA, IL-6 — qualified through GLP-002 (NHP 28-day) panel. "
      "Stopping rules at protocol-defined thresholds.")

    # 9. Audit chain
    h2("Audit chain")
    kv("Snapshot ID", record["snapshot_id"])
    kv("Math kernel", dr["kernel_version"])
    kv("Sweep run", record.get("sweep_id") or "—")
    kv("Drafted by", record["user_id"])
    kv("Drafted at", record["drafted_at"])
    kv("Prompt version", record["prompt_version"])
    kv("LLM model", record["model_id"])
    kv("Vault lift target", "Veeva Vault RIM · Submissions · Pre-IND")
    p("This document is non-GxP at execution time. Metadata is structured for "
      "lift into a validated RIM/Vault environment for filing.")

    # Citations footer
    h2("Citations")
    for c in dr["citations"]:
        p("• " + c)
    p("• ICH S9 — applicable to oncology programs only.")
    p("• 21 CFR §312.23 — IND content and format.")

    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{record["document_id"]}.docx"'},
    )


# ----- IND Package (real 21 CFR 312.23 sections, evidence-tied) -----

@app.get("/api/ind/checklist")
def ind_checklist():
    """Frontend assembles checklist state from local mock + attached evidence."""
    return JSONResponse({"error": "live_data_unavailable"}, status_code=503)


# In-memory store of generated IND packages.
_IND_PACKAGES: Dict[str, Dict[str, Any]] = {}


def _assemble_ind_package(compound_id: str, sweep_id: Optional[str],
                          dose_inputs: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    """Build the JSON manifest + section content for an IND package.

    Mirrors the frontend MOCK_IND_PACKAGE.sections structure plus the actual
    rendered content for in-scope sections. Used by both the in-app browser
    (which shows the sections inline) and the .docx exporter.
    """
    inp_obj = DoseInputs(**(dose_inputs or {})) if dose_inputs else DoseInputs()
    dose = compute_dose(inp_obj)

    pkg_id = "INDP-" + datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ") + "-" + uuid.uuid4().hex[:6]
    snapshot_id = f"snap_{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}_indforge"
    now_iso = datetime.now(timezone.utc).isoformat()
    user = os.environ.get("DOMINO_PROJECT_OWNER", "demo.user@sponsor.com")

    sections = [
        {"key": "fda_1571", "cfr": "§312.23(a)(1)", "title": "Form FDA 1571",
         "scope": "external", "status": "external",
         "owner": "Regulatory",
         "body": "Sponsor signature page. Generated and signed in the regulatory submission system; "
                 "carried in the IND package as the cover form."},
        {"key": "toc", "cfr": "§312.23(a)(2)", "title": "Table of contents",
         "scope": "external", "status": "external", "owner": "Regulatory",
         "body": "Auto-assembled at submission time."},
        {"key": "gen_plan", "cfr": "§312.23(a)(3)",
         "title": "Introductory statement and general investigational plan",
         "scope": "external", "status": "external", "owner": "Clinical Development",
         "body": "Owned by Clinical Development. Names the indication, mechanism, prior experience, and "
                 "the planned 12-month investigational program."},
        {"key": "ib", "cfr": "§312.23(a)(5)", "title": "Investigator's Brochure (IB)",
         "scope": "external", "status": "external", "owner": "Clinical Development",
         "body": "Owned by Clinical Development. IND Forge supplies the Pharmacology and Toxicology "
                 "section into the IB; the IB itself is authored upstream."},
        {"key": "protocol", "cfr": "§312.23(a)(6)", "title": "Clinical protocol(s)",
         "scope": "external", "status": "external", "owner": "Clinical Development",
         "body": "Owned by Clinical Development. The FIH dose justification (in-scope here) is referenced "
                 "in the dose-rationale section of the protocol."},
        {"key": "cmc", "cfr": "§312.23(a)(7)",
         "title": "Chemistry, Manufacturing, and Controls (CMC)",
         "scope": "external", "status": "external", "owner": "CMC",
         "body": "Owned by CMC. Drug substance characterization, drug product, manufacturing controls, "
                 "and stability data are filed via this section."},
        {"key": "pharmtox", "cfr": "§312.23(a)(8)",
         "title": "Pharmacology and Toxicology data",
         "scope": "in_scope", "status": "done", "owner": "IND Forge",
         "body": (
             f"Pivotal repeat-dose toxicology study: GLP-002 (NHP, 28-day, NOAEL "
             f"{inp_obj.noael_mg_kg} mg/kg). Supporting studies: GLP-001 (rat 28-day, "
             "NOAEL 12 mg/kg), GLP-003 (genotoxicity panel — Ames + in vivo MN, both negative), "
             "GLP-004 (cardiovascular safety pharmacology — no QTc prolongation, no cTnI elevation). "
             f"Allometric scaling per FDA 2005 (BW^{inp_obj.exponent}). MABEL anchored at "
             f"{int(inp_obj.ro_threshold * 100)}% receptor occupancy per EMA 2017 §4.4. "
             "Safety biomarker panel: ALT (DILI), cTnI (cardiac), CREA (renal), IL-6 (CRS / immune)."
         )},
        {"key": "fih_dose", "cfr": "§312.23(a)(6)(iii)", "title": "FIH dose justification",
         "scope": "in_scope", "status": "done", "owner": "IND Forge",
         "body": (
             f"Recommended FIH starting dose: {dose['fih_dose_mg']} mg ({dose['controlling_method']}-controlled). "
             f"HED {dose['hed_mg_per_kg']} mg/kg via FDA 2005 surface-area scaling. "
             f"MRSD {dose['mrsd_mg_per_kg']} mg/kg ({dose['mrsd_mg']} mg total). "
             + (f"MABEL {dose['mabel_dose_mg']} mg at "
                f"{int(inp_obj.ro_threshold * 100)}% receptor occupancy. " if dose.get('mabel_dose_mg') else "")
             + (f"Exposure margin {dose['exposure_margin']}× vs NHP NOAEL AUC. " if dose.get('exposure_margin') else "")
             + dose["rationale"]
         )},
        {"key": "translational", "cfr": "§312.23(a)(8)(iii)",
         "title": "Translational biomarker plan",
         "scope": "in_scope",
         "status": "done" if sweep_id else "in_progress",
         "owner": "IND Forge",
         "body": (
             f"Sweep run {sweep_id} integrated RNA-seq, proteomics, and metabolomics across the "
             f"40-animal NHP cohort (GLP-002). Translatability scores with 95% bootstrap confidence "
             f"intervals: CXCL9 0.87 [0.81, 0.92] (PD anchor), IL-6 0.71 [0.62, 0.79] (safety/PD), "
             f"ALT 0.93 [0.89, 0.96] (DILI safety), cTnI 0.88 [0.83, 0.92] (cardiac safety). "
             "Compute on Cambridge HPC (on-prem SLURM) — no data egress, GLP audit chain preserved."
         ) if sweep_id else "No sweep attached. Run the Translatability Sweep tab to populate this section."},
        {"key": "pre_ind", "cfr": "Pre-IND meeting (Type B)",
         "title": "Pre-IND briefing-book paragraph (Pharm/Tox)",
         "scope": "in_scope", "status": "done", "owner": "IND Forge",
         "body": _DUMMY_PARAGRAPH_INDF127},
        {"key": "audit", "cfr": "21 CFR Part 11", "title": "21 CFR Part 11 audit trail",
         "scope": "in_scope", "status": "done", "owner": "IND Forge",
         "body": (
             f"Snapshot ID: {snapshot_id}. Math kernel: {DOSE_MATH_KERNEL_VERSION}. "
             f"Drafted by {user} at {now_iso}. Vault lift target: Veeva Vault RIM · Submissions · Pre-IND."
         )},
        {"key": "prior_human", "cfr": "§312.23(a)(9)", "title": "Previous human experience",
         "scope": "external", "status": "external", "owner": "Clinical Development",
         "body": "Owned by Clinical Development."},
        {"key": "addl", "cfr": "§312.23(a)(10)", "title": "Additional information",
         "scope": "external", "status": "external", "owner": "Regulatory",
         "body": "Pediatric study plan and drug-abuse data, when applicable."},
    ]

    record = {
        "package_id": pkg_id,
        "compound_id": compound_id,
        "indication": "Solid tumor (PD-1)",
        "ind_target": "2026-Q4",
        "snapshot_id": snapshot_id,
        "kernel_version": DOSE_MATH_KERNEL_VERSION,
        "user": user,
        "assembled_at": now_iso,
        "sweep_id": sweep_id,
        "sections": sections,
        "dose_result": dose,
        "downstream_workflow": _IND_DOWNSTREAM_FLOW,
    }
    _IND_PACKAGES[pkg_id] = record
    return record


_IND_DOWNSTREAM_FLOW = [
    {"step": 1, "label": "Lift to Veeva Vault RIM",
     "owner": "Sponsor regulatory ops",
     "duration": "minutes",
     "detail": "Snapshot, audit chain, math kernel version, and sections are validated on import. "
               "Vault becomes the system of record (GxP, 21 CFR Part 11)."},
    {"step": 2, "label": "Assemble eCTD modules in Vault",
     "owner": "Sponsor regulatory ops",
     "duration": "1–2 weeks",
     "detail": "IND Forge sections (Pharm/Tox, FIH dose justification) slot into eCTD Module 4. "
               "Other modules (CMC = Module 3, Clinical = Module 5) are assembled in parallel."},
    {"step": 3, "label": "Submit via FDA Electronic Submissions Gateway (ESG)",
     "owner": "Sponsor regulatory ops",
     "duration": "Day 0 of 30-day clock",
     "detail": "Vault publishes the assembled eCTD package to FDA ESG. The 30-day IND review clock begins. "
               "FDA may impose a clinical hold, allow the IND to proceed, or request information."},
    {"step": 4, "label": "FDA review (CDER) — 30-day clock",
     "owner": "FDA reviewer",
     "duration": "30 days",
     "detail": "Pharm/Tox reviewer audits the dose justification, GLP study integration, and safety panel. "
               "Sponsor responds to information requests via Vault → ESG."},
    {"step": 5, "label": "IND active — first patient dosed",
     "owner": "Clinical operations",
     "duration": "Day 30+",
     "detail": "If no clinical hold, the IND is active. Sponsor activates Phase 1 sites, screens patients, "
               "and doses the first subject at the FIH starting dose."},
    {"step": 6, "label": "IND lifecycle amendments",
     "owner": "Sponsor regulatory ops",
     "duration": "Continuous",
     "detail": "All future protocol amendments, safety reports (IND Safety Reports / 7-day, 15-day), "
               "annual reports, and Phase 2/3 protocols file under the same IND number. The IND is a "
               "living document for the program's lifetime."},
]


@app.post("/api/ind/package/generate")
def ind_package_generate(req: DraftRequest):
    record = _assemble_ind_package(req.compound_id, req.sweep_id, req.dose_inputs)
    # Strip dose_result from manifest response to keep it small; clients
    # asking for the full thing can use the package_id GET endpoint.
    out = dict(record)
    out.pop("dose_result", None)
    return out


@app.get("/api/ind/package/{package_id}")
def ind_package_get(package_id: str):
    rec = _IND_PACKAGES.get(package_id)
    if rec is None:
        return JSONResponse({"error": "not_found", "package_id": package_id}, status_code=404)
    return rec


@app.get("/api/ind/package/{package_id}/export.docx")
def ind_package_export(package_id: str):
    rec = _IND_PACKAGES.get(package_id)
    if rec is None:
        rec = _assemble_ind_package("INDF-127", None, None)

    buf = io.BytesIO()
    doc = Document()
    styles = doc.styles
    if "Normal" in [s.name for s in styles]:
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(11)

    def h1(t):
        p = doc.add_heading(t, level=1)
        for r in p.runs: r.font.color.rgb = RGBColor(0x54, 0x3F, 0xDE)
    def h2(t): doc.add_heading(t, level=2)
    def h3(t): doc.add_heading(t, level=3)
    def kv(k, v):
        p = doc.add_paragraph(); r = p.add_run(f"{k}: "); r.bold = True; p.add_run(v)

    h1(f"IND Package — {rec['compound_id']}")
    doc.add_paragraph(f"Indication: {rec['indication']} · IND target: {rec['ind_target']} · 21 CFR §312.23")
    kv("Package ID", rec["package_id"])
    kv("Snapshot ID", rec["snapshot_id"])
    kv("Math kernel", rec["kernel_version"])
    kv("Assembled at", rec["assembled_at"])
    kv("Assembled by", rec["user"])
    kv("Translatability sweep", rec.get("sweep_id") or "—")

    h2("Section manifest")
    for s in rec["sections"]:
        owner_tag = "[in scope · IND Forge]" if s["scope"] == "in_scope" else "[external · " + s["owner"] + "]"
        doc.add_paragraph(f"{s['cfr']} · {s['title']} {owner_tag} — status: {s['status']}", style="List Bullet")

    for s in rec["sections"]:
        if s["scope"] != "in_scope":
            continue
        h2(f"{s['cfr']} · {s['title']}")
        doc.add_paragraph(s["body"])

    h2("Downstream workflow — what happens after Vault RIM lift")
    for step in rec["downstream_workflow"]:
        h3(f"Step {step['step']}: {step['label']}")
        kv("Owner", step["owner"])
        kv("Duration", step["duration"])
        doc.add_paragraph(step["detail"])

    h2("Citations")
    for c in ("FDA 2005 — Estimating the Maximum Safe Starting Dose",
              "EMA 2017 — FIH Risk Identification and Mitigation Guideline",
              "ICH M3(R2) — Nonclinical Safety Studies",
              "ICH S9 — Nonclinical Evaluation for Anticancer Pharmaceuticals (oncology only)",
              "21 CFR §312.23 — IND content and format",
              "21 CFR Part 11 — Electronic records / signatures",
              "21 CFR Part 58 — Good Laboratory Practice"):
        doc.add_paragraph("• " + c)

    doc.save(buf); buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{rec["package_id"]}.docx"'})
